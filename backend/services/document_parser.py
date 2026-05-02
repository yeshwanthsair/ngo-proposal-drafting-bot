"""
Document parsing service.
Handles PDF, DOCX, and TXT file parsing and text chunking.
"""
import os
import logging
from typing import List
from pathlib import Path

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def parse_txt(file_path: str) -> str:
    """Parse plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_pdf(file_path: str) -> str:
    """Parse PDF file using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        return text
    except Exception as e:
        logger.error(f"Error parsing PDF {file_path}: {e}")
        raise


def parse_docx(file_path: str) -> str:
    """Parse DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error parsing DOCX {file_path}: {e}")
        raise


def parse_document(file_path: str) -> str:
    """
    Parse a document based on its file extension.
    Supports: .txt, .pdf, .docx
    """
    ext = Path(file_path).suffix.lower()
    parsers = {
        ".txt": parse_txt,
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".doc": parse_docx,
    }

    if ext not in parsers:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(parsers.keys())}")

    logger.info(f"Parsing {ext} file: {file_path}")
    return parsers[ext](file_path)


def chunk_text(
    text: str,
    filename: str,
    chunk_size: int = 800,
    chunk_overlap: int = 150
) -> List[Document]:
    """
    Split text into overlapping chunks for vector storage.

    Args:
        text: Raw document text
        filename: Source filename (used as metadata)
        chunk_size: Max characters per chunk
        chunk_overlap: Overlap between consecutive chunks

    Returns:
        List of LangChain Document objects with metadata
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )

    chunks = splitter.split_text(text)

    documents = [
        Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "chunk_index": i,
                "total_chunks": len(chunks),
            }
        )
        for i, chunk in enumerate(chunks)
    ]

    logger.info(f"Created {len(documents)} chunks from '{filename}'")
    return documents


def parse_and_chunk(file_path: str, chunk_size: int = 800, chunk_overlap: int = 150) -> List[Document]:
    """
    Full pipeline: parse a file and return chunked Documents.
    """
    filename = Path(file_path).name
    raw_text = parse_document(file_path)

    if not raw_text.strip():
        raise ValueError(f"No text content found in '{filename}'")

    return chunk_text(raw_text, filename, chunk_size, chunk_overlap)
